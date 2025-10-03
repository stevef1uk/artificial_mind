#!/usr/bin/env python3
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    sys.stderr.write("Missing dependency: python-docx. Install with: pip install python-docx\n")
    raise

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

def read_file(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return ""

def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h

def add_paragraph(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    return p

def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)

def section_divider(doc: Document):
    doc.add_page_break()

def build_document():
    doc = Document()

    # Title
    title = doc.add_heading('Artificial Mind Solution Architecture', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%SZ')} (UTC)")

    # Inputs
    summary_md = read_file(os.path.join(ROOT, 'SUMMARY.md'))
    arch_md = read_file(os.path.join(ROOT, 'ARCHITECTURE.md'))
    sys_overview_md = read_file(os.path.join(ROOT, 'SYSTEM_OVERVIEW.md'))

    # Context & High-Level View
    section_divider(doc)
    add_heading(doc, 'Context and High-Level View', level=1)
    add_paragraph(doc, 'This document summarizes the artificial mind solution architecture, contextualizes the system, and presents a high-level view of the major subsystems and their interactions.', False)
    add_heading(doc, 'High-Level Architecture', level=2)
    if sys_overview_md:
        add_paragraph(doc, 'Overview (from SYSTEM_OVERVIEW.md):', bold=True)
        add_paragraph(doc, sys_overview_md[:2000] + ('...' if len(sys_overview_md) > 2000 else ''))
    else:
        add_bullets(doc, [
            'Cognition & Policy: FSM, Self-Model & Goal Manager, Principles Server',
            'Planning & Execution: HDN API, Planner/Evaluator, Intelligent Executor, Code Generator',
            'Data & Infra: Redis, Qdrant, Neo4j, Docker',
            'Eventing: NATS canonical event bus',
        ])

    # Component Deep Dives (next level down)
    section_divider(doc)
    add_heading(doc, 'Component Deep Dives', level=1)

    add_heading(doc, 'FSM Engine (Control Layer)', level=2)
    add_bullets(doc, [
        'State-driven cognition: perception → learning → planning → evaluation → execution',
        'Reasoning layer: belief querying, forward-chaining inference, curiosity goals, explanation traces',
        'Integration: Principles gates, knowledge growth, HDN delegation, NATS events',
    ])

    add_heading(doc, 'HDN (Planning & Execution)', level=2)
    add_bullets(doc, [
        'Intelligent code generation and validation (LLM + Docker)',
        'Capability learning, caching, and dynamic action creation',
        'Planner/Evaluator integration; workflow orchestration; file/artifact management',
    ])

    add_heading(doc, 'Self-Model & Goal Manager (Motivation)', level=2)
    add_bullets(doc, [
        'Goals, beliefs, episodic history persisted in Redis',
        'Policy layer prioritizes goals; emits agi.goal.* over NATS',
        'Learning loop updates confidence, priorities, and performance metrics',
    ])

    add_heading(doc, 'Principles Server (Ethics & Safety)', level=2)
    add_bullets(doc, [
        'JSON-based rules; dynamic reload; context-aware checks',
        'Pre-exec gates for tools/actions; audit trails and denials with reasons',
        'Layered with LLM safety categorization and Docker sandboxing',
    ])

    add_heading(doc, 'Memory & Knowledge', level=2)
    add_bullets(doc, [
        'Working Memory (Redis): ephemeral state, capabilities, workflow artifacts',
        'Episodic Memory (Qdrant): vector search over episodes for retrieval-augmented reasoning',
        'Semantic Knowledge (Neo4j): domain concepts, relations, constraints, safety principles',
    ])

    # Technical Architecture
    section_divider(doc)
    add_heading(doc, 'Technical Architecture', level=1)
    add_bullets(doc, [
        'APIs: RESTful services (HDN 8081, Principles 8080, Monitor 8082)',
        'Event Bus: NATS (agi.events.*) for canonical event envelopes',
        'Execution Sandbox: Docker with resource limits, timeouts, and isolation',
        'Persistence: Redis (state/cache), Qdrant (episodes), Neo4j (knowledge)',
        'Observability: Monitor UI, health checks, metrics, workflow and artifact views',
        'Deployment: Docker/K3s manifests, cronjobs for scheduled tasks',
        'Security: Principles gating, content safety, tool metrics and audit logging',
    ])

    # Viability Summary (from SUMMARY.md)
    section_divider(doc)
    add_heading(doc, 'Viability Summary', level=1)
    if summary_md:
        add_paragraph(doc, 'Highlights (from SUMMARY.md):', bold=True)
        add_paragraph(doc, summary_md[:3000] + ('...' if len(summary_md) > 3000 else ''))
    else:
        add_bullets(doc, [
            'Self-improving loop of learning, caching, and capability reuse',
            'Strong safety posture with multi-layer defenses',
            'Scalable, modular, event-driven design with robust monitoring',
        ])

    return doc

def main():
    out_dir = os.path.join(ROOT, 'docs')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, 'AGI_Solution_Architecture.docx')
    doc = build_document()
    doc.save(out_path)
    print(out_path)

if __name__ == '__main__':
    main()


